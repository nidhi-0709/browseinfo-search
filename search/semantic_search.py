import os
import json
import time
import faiss
import numpy as np
from datetime import datetime
from sentence_transformers import SentenceTransformer
from openpyxl import Workbook
from docx import Document

# =====================================================
# Performance Settings
# =====================================================

faiss.omp_set_num_threads(8)

# =====================================================
# Load Embedding Model
# =====================================================

print("Loading AI Model...")
model = SentenceTransformer("all-MiniLM-L6-v2")

# =====================================================
# Paths
# =====================================================

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

INDEX_FILE = os.path.join(BASE_DIR, "odoo_modules_index.faiss")
JSON_FILE = os.path.join(BASE_DIR, "merged_modules.json")

# =====================================================
# Load FAISS Index & JSON
# =====================================================

print("Loading FAISS Index...")
index = faiss.read_index(INDEX_FILE)

print("Loading Module Data...")

with open(JSON_FILE, "r", encoding="utf-8") as f:
    modules = json.load(f)

print(f"{len(modules)} modules loaded successfully.")

# =====================================================
# Report Files
# =====================================================

timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

excel_file = os.path.join(
    BASE_DIR,
    f"Odoo_AI_Report_{timestamp}.xlsx"
)

word_file = os.path.join(
    BASE_DIR,
    f"Odoo_AI_Report_{timestamp}.docx"
)

wb = Workbook()
ws = wb.active
ws.title = "Recommendations"

ws.append([
    "Requirement",
    "Rank",
    "Module",
    "Source",
    "Category",
    "Score",
    "Summary",
    "Features",
    "URL"
])

doc = Document()
doc.add_heading("Odoo AI Recommendation Report", level=1)
doc.add_paragraph(f"Generated on : {datetime.now()}")

# =====================================================
# Semantic Search
# =====================================================

def search_modules(requirement, top_k=5):

    embedding = model.encode(
        requirement,
        convert_to_numpy=True,
        normalize_embeddings=True,
        show_progress_bar=False
    )

    embedding = embedding.reshape(1, -1).astype(np.float32)

    scores, ids = index.search(
        embedding,
        top_k
    )

    results = []

    for score, idx in zip(scores[0], ids[0]):

        if idx == -1:
            continue

        module = modules[idx]

        results.append({

            "Title": module.get("Title", ""),

            "Source": module.get("Source", ""),

            "Category": module.get("Category", ""),

            "Summary": module.get("Summary", ""),

            "Features": module.get("Features", ""),

            "URL": module.get("URL", ""),

            "Score": float(score)

        })

    return results

# =====================================================
# Main
# =====================================================

print("\n" + "="*70)
print("         ODOO AI MODULE RECOMMENDER")
print("="*70)

print("\nType 'exit' to save reports and quit.\n")

while True:

    requirement = input("Enter Client Requirement:\n> ").strip()

    if requirement.lower() in ["exit", "quit", "q"]:

        wb.save(excel_file)
        doc.save(word_file)

        print("\nReports Generated Successfully.")
        print(f"\nExcel Report : {excel_file}")
        print(f"Word Report  : {word_file}")

        print("\nThank you for using Odoo AI Module Recommender.")
        break

    start = time.perf_counter()

    results = search_modules(
        requirement,
        top_k=5
    )

    elapsed = time.perf_counter() - start

    print(f"\nSearch completed in {elapsed:.3f} seconds\n")

    print("="*90)
    print("TOP RECOMMENDED MODULES")
    print("="*90)

    doc.add_heading("Client Requirement", level=2)
    doc.add_paragraph(requirement)

    doc.add_heading("Recommended Modules", level=2)

    for i, module in enumerate(results, start=1):

        features = module["Features"]

        if isinstance(features, list):
            features = ", ".join(features)

        print(f"\n{i}. {module['Title']}")
        print(f"Source   : {module['Source']}")
        print(f"Category : {module['Category']}")
        print(f"Score    : {module['Score']:.4f}")

        # Excel

        ws.append([
            requirement,
            i,
            module["Title"],
            module["Source"],
            module["Category"],
            round(module["Score"],4),
            module["Summary"],
            str(features),
            module["URL"]
        ])

        # Word

        report = f"""
{i}. {module['Title']}

Source : {module['Source']}
Category : {module['Category']}
Score : {module['Score']:.4f}

Summary:
{module['Summary']}

Features:
{features}

URL:
{module['URL']}

"""

        doc.add_paragraph(report)

    doc.add_page_break()

    print("\n" + "-"*90 + "\n")